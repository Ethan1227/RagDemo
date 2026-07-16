"""起诉状 Word 导出（python-docx）。

将起诉状文本（Markdown）渲染为标准法律文书 Word：
标题黑体居中、正文宋体小四、段落首行缩进、落款右对齐。
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

_TITLE = "民事起诉状"
# 落款关键词：以此开头的行右对齐
_RIGHT_ALIGN_PREFIX = ("此致", "具状人", "起诉人", "年", "20")


def _set_font(run, name: str = "宋体", size: int = 14, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    # 中文字体需单独设置 eastAsia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _clean_markdown(line: str) -> tuple[str, int]:
    """去除 Markdown 标记，返回 (纯文本, 标题级别 0=正文)。"""
    heading = re.match(r"^(#{1,6})\s+(.*)", line)
    if heading:
        return heading.group(2).strip(), len(heading.group(1))
    # 去除加粗/列表符号
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
    text = re.sub(r"^[-*]\s+", "", text)
    return text.strip(), 0


def render_docx(content: str, title: str = _TITLE) -> bytes:
    """将起诉状内容渲染为 Word，返回字节。"""
    doc = Document()

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    _set_font(title_run, name="黑体", size=22, bold=True)

    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        # 跳过与标题重复的首行
        if line.strip() in (title, f"# {title}", _TITLE):
            continue

        text, level = _clean_markdown(line)
        if not text:
            continue

        para = doc.add_paragraph()
        if level and level <= 2:
            run = para.add_run(text)
            _set_font(run, name="黑体", size=15, bold=True)
        else:
            run = para.add_run(text)
            _set_font(run, name="宋体", size=14)
            if text.startswith(_RIGHT_ALIGN_PREFIX):
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                para.paragraph_format.first_line_indent = Pt(28)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
