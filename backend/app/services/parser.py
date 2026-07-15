"""文档解析与切块服务。

阶段 2 支持：PDF(pypdf) / Word(.docx, python-docx) / Markdown / TXT。
图片类文档（OCR）在阶段 3 接入 PaddleOCR 后支持。

切块策略：按句子/段落边界贪心装包，块间保留 chunk_overlap 字符重叠；
单句超长时按 chunk_size 硬切。
"""
from __future__ import annotations

import io
import re
from pathlib import Path

from loguru import logger

SUPPORTED_TYPES = {"pdf", "docx", "md", "markdown", "txt"}

# 句末标点（含换行），用于把文本切成"原子"句段
_SENTENCE_RE = re.compile(r"[^。！？!?\n]*[。！？!?\n]|[^。！？!?\n]+$")


class UnsupportedFileTypeError(ValueError):
    """不支持的文件类型。"""


def detect_file_type(filename: str) -> str:
    """从文件名推断类型，不支持时显式抛错。"""
    suffix = Path(filename).suffix.lower().lstrip(".")
    if suffix == "markdown":
        suffix = "md"
    if suffix not in SUPPORTED_TYPES:
        raise UnsupportedFileTypeError(
            f"不支持的文件类型：.{suffix}（阶段2支持 pdf/docx/md/txt，图片将在阶段3支持）"
        )
    return suffix


def extract_text(filename: str, data: bytes) -> str:
    """按类型提取纯文本。"""
    file_type = detect_file_type(filename)
    if file_type == "pdf":
        return _extract_pdf(data)
    if file_type == "docx":
        return _extract_docx(data)
    # md / txt：按 utf-8 解码，失败回退 gbk
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("gbk", errors="replace")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages)
    logger.debug("PDF 提取完成：{} 页，{} 字符", len(pages), len(text))
    return text


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    # 表格内容也提取（法律文书常含表格）
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    text = "\n".join(parts)
    logger.debug("DOCX 提取完成：{} 字符", len(text))
    return text


def _split_sentences(text: str) -> list[str]:
    """把文本切成带结尾标点的句段（原子单位）。"""
    sentences = [s for s in _SENTENCE_RE.findall(text) if s.strip()]
    return sentences


def split_text(text: str, chunk_size: int = 512, chunk_overlap: int = 50) -> list[str]:
    """按句子边界切块，块间保留重叠。

    - chunk_overlap 必须小于 chunk_size（否则无法前进），违反时显式抛错
    - 空文本返回空列表
    """
    if chunk_overlap >= chunk_size:
        raise ValueError(
            f"chunk_overlap({chunk_overlap}) 必须小于 chunk_size({chunk_size})"
        )
    text = re.sub(r"[ \t]+", " ", text).strip()
    if not text:
        return []

    # 超长句子先硬切，保证每个原子 <= chunk_size
    atoms: list[str] = []
    for sentence in _split_sentences(text):
        while len(sentence) > chunk_size:
            atoms.append(sentence[:chunk_size])
            sentence = sentence[chunk_size:]
        if sentence:
            atoms.append(sentence)

    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    for atom in atoms:
        if current_len + len(atom) > chunk_size and current:
            chunks.append("".join(current).strip())
            # 从当前块尾部取 ~chunk_overlap 字符的完整句段作为下一块开头
            overlap: list[str] = []
            overlap_len = 0
            for s in reversed(current):
                if overlap_len + len(s) > chunk_overlap:
                    break
                overlap.insert(0, s)
                overlap_len += len(s)
            current = overlap
            current_len = overlap_len
        current.append(atom)
        current_len += len(atom)
    if current:
        chunks.append("".join(current).strip())

    return [c for c in chunks if c]
